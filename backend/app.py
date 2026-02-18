# app.py - Backend Flask complet avec corrections
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
import pandas as pd
import numpy as np
import os
import json
from datetime import datetime
import uuid
import re
import xml.etree.ElementTree as ET
import logging
from config import get_config, get_message
from auth import auth_bp, token_required

from zipfile import ZipFile
from io import BytesIO

from assistant import DataAssistant
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.DEBUG)

app = Flask(__name__)

SECRET_KEY = os.getenv("SECRET_KEY")
app.config['SECRET_KEY'] = SECRET_KEY
configclass = get_config()
app.config.from_object(configclass)
CORS(app,
     resources={r"/api/*": {"origins": "*"}},
     supports_credentials=True,
     allow_headers=["Content-Type", "Authorization"],
     methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"]
     )

app.register_blueprint(auth_bp, url_prefix="/api")

if os.environ.get('RENDER'):
    UPLOAD_FOLDER = '/tmp/uploads'
    CLEANED_FOLDER = '/tmp/cleaned'
else:
    UPLOAD_FOLDER = 'uploads'
    CLEANED_FOLDER = 'cleaned'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CLEANED_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CLEANED_FOLDER'] = CLEANED_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

sessions_db = {}


# -------------------- UTILITAIRES --------------------

def convert_to_serializable(obj):
    if isinstance(obj, np.integer):
        return int(obj)
    elif isinstance(obj, np.floating):
        if np.isnan(obj) or np.isinf(obj):
            return None
        return float(obj)
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    elif isinstance(obj, pd.Timestamp):
        return obj.isoformat()
    elif isinstance(obj, dict):
        return {k: convert_to_serializable(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [convert_to_serializable(item) for item in obj]
    try:
        if pd.isna(obj):
            return None
    except Exception:
        pass
    return obj


def allowed_file(filename):
    return (
        '.' in filename and
        filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']
    )


def load_file(filepath, file_extension):
    try:
        if file_extension == 'csv':
            encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
            separators = [',', ';', '\t', '|']
            for encoding in encodings:
                for sep in separators:
                    try:
                        df = pd.read_csv(filepath, encoding=encoding, sep=sep, on_bad_lines='skip')
                        if len(df.columns) > 1:
                            df.replace(r'^\s*$', np.nan, regex=True, inplace=True)
                            return df
                    except Exception:
                        continue
            df = pd.read_csv(filepath, encoding='utf-8', sep=',', on_bad_lines='skip')
            df.replace(r'^\s*$', np.nan, regex=True, inplace=True)
            return df
        elif file_extension in ['xlsx', 'xls']:
            df = pd.read_excel(filepath)
            df.replace(r'^\s*$', np.nan, regex=True, inplace=True)
            return df
        elif file_extension == 'json':
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            df = pd.DataFrame(data)
            df.replace(r'^\s*$', np.nan, regex=True, inplace=True)
            return df
        elif file_extension == 'xml':
            tree = ET.parse(filepath)
            root = tree.getroot()
            data = []
            for child in root:
                row = {elem.tag: elem.text for elem in child}
                data.append(row)
            df = pd.DataFrame(data)
            df.replace(r'^\s*$', np.nan, regex=True, inplace=True)
            return df
        else:
            raise ValueError(f"Extension {file_extension} non supportée")
    except Exception as e:
        logging.error(f"[load_file ERROR] {str(e)}")
        raise Exception(f"Impossible de charger le fichier: {str(e)}")


def save_file(df, filepath, file_extension):
    try:
        if file_extension == 'csv':
            df.to_csv(filepath, index=False, encoding='utf-8')
        elif file_extension in ['xlsx', 'xls']:
            df.to_excel(filepath, index=False)
        elif file_extension == 'json':
            df.to_json(filepath, orient='records', indent=2)
        elif file_extension == 'xml':
            root = ET.Element('data')
            for _, row in df.iterrows():
                record = ET.SubElement(root, 'record')
                for col in df.columns:
                    record_elem = ET.SubElement(record, str(col))
                    val = row[col]
                    record_elem.text = '' if pd.isna(val) else str(val)
            tree = ET.ElementTree(root)
            tree.write(filepath, encoding='utf-8', xml_declaration=True)
    except Exception as e:
        logging.error(f"[save_file ERROR] {str(e)}")
        raise Exception(f"Erreur lors de la sauvegarde: {str(e)}")


# -------------------- DÉTECTION COLONNES MIXTES --------------------

def _is_truly_mixed_column(series: pd.Series) -> bool:
    """
    Retourne True UNIQUEMENT si la colonne contient À LA FOIS :
      - au moins une valeur numérique valide
      - au moins une valeur texte non-numérique non-nulle

    Exemples :
      [3, 1, 'HURLEY', 2]       → True   (mixte réelle)
      ['PUTNAM', 'BERKELEY']    → False  (texte pur → NE PAS TOUCHER)
      [1.5, 2.0, np.nan]        → False  (numérique avec manquants seulement)
      ['--', 'na', 1, 2]        → True   (pseudo-nulls texte + numériques)
      [np.nan, np.nan]          → False  (tout null)
    """
    non_null = series.dropna()
    if len(non_null) == 0:
        return False

    numeric_converted = pd.to_numeric(non_null, errors='coerce')
    n_numeric = int(numeric_converted.notna().sum())   # valeurs vraiment numériques
    n_text = int(numeric_converted.isna().sum())       # valeurs non convertibles = texte

    # Les DEUX catégories doivent être présentes pour qualifier de "mixte"
    return n_numeric > 0 and n_text > 0


# -------------------- ANALYSE DES DONNÉES --------------------

class DataAnalyzer:

    @staticmethod
    def detect_column_types(df):
        types = {}
        for col in df.columns:
            try:
                if pd.api.types.is_numeric_dtype(df[col]):
                    types[col] = 'numeric'
                elif pd.api.types.is_datetime64_any_dtype(df[col]):
                    types[col] = 'datetime'
                else:
                    sample = df[col].dropna().head(100)
                    detected = False
                    for fmt in app.config['DATE_FORMATS']:
                        try:
                            parsed = pd.to_datetime(sample, format=fmt, errors='coerce')
                            if parsed.notna().sum() > 0:
                                types[col] = 'date_string'
                                detected = True
                                break
                        except Exception:
                            continue
                    if not detected:
                        types[col] = 'text'
            except Exception:
                types[col] = 'text'
        return types

    @staticmethod
    def detect_mixed_columns(df):
        """
        Retourne uniquement les colonnes avec un vrai mélange num/texte.
        Colonnes purement textuelles (ST_NAME, OWN_OCCUPIED…) sont ignorées.
        """
        return [col for col in df.columns if _is_truly_mixed_column(df[col])]

    @staticmethod
    def analyze_missing_values(df):
        missing = {}
        for col in df.columns:
            count = df[col].isna().sum()
            if count > 0:
                missing[col] = {
                    'count': int(count),
                    'percentage': round(float(count / len(df)) * 100, 2)
                }
        return missing

    @staticmethod
    def detect_duplicates(df, uniqueness_threshold=0.95):
        exact_count = int(df.duplicated().sum())
        candidate_cols = []
        for col in df.columns:
            uniqueness_ratio = df[col].nunique(dropna=True) / len(df)
            if uniqueness_ratio < uniqueness_threshold:
                candidate_cols.append(col)
        structural_count = int(df.duplicated(subset=candidate_cols).sum()) if candidate_cols else 0
        return {
            "exact_duplicates": exact_count,
            "structural_duplicates": structural_count,
            "used_columns": candidate_cols
        }

    @staticmethod
    def detect_outliers(df, column_types):
        outliers = {}
        for col in df.columns:
            if column_types.get(col) == 'numeric':
                numeric_col = pd.to_numeric(df[col], errors='coerce')
                Q1 = numeric_col.quantile(0.25)
                Q3 = numeric_col.quantile(0.75)
                IQR = Q3 - Q1
                lower = Q1 - 1.5 * IQR
                upper = Q3 + 1.5 * IQR
                count = int(((numeric_col < lower) | (numeric_col > upper)).sum())
                if count > 0:
                    outliers[col] = count
        return outliers

    @staticmethod
    def analyze_text_issues(df, column_types):
        issues = {}
        emoji_pattern = re.compile(
            "["
            u"\U0001F600-\U0001F64F"
            u"\U0001F300-\U0001F5FF"
            u"\U0001F680-\U0001F6FF"
            u"\U0001F1E0-\U0001F1FF"
            "]+", flags=re.UNICODE
        )
        for col in df.columns:
            if column_types.get(col) in ['text', 'date_string']:
                col_issues = {}
                emojis = int(df[col].apply(lambda x: bool(emoji_pattern.search(str(x)))).sum())
                if emojis > 0:
                    col_issues['emojis'] = emojis
                extra_spaces = int(df[col].apply(
                    lambda x: '  ' in str(x) or str(x) != str(x).strip()
                ).sum())
                if extra_spaces > 0:
                    col_issues['spaces'] = extra_spaces
                special_chars = int(df[col].apply(
                    lambda x: bool(re.search(r'[^a-zA-Z0-9\s\-_.,@]', str(x)))
                ).sum())
                if special_chars > 0:
                    col_issues['specialChars'] = special_chars
                if df[col].nunique() < len(df) * 0.5:
                    values_lower = df[col].apply(lambda x: str(x).lower())
                    inconsistent = len(df[col].unique()) - len(values_lower.unique())
                    if inconsistent > 0:
                        col_issues['inconsistentCase'] = int(inconsistent)
                if col_issues:
                    issues[col] = col_issues
        return issues

    @staticmethod
    def detect_date_formats(df, column_types):
        date_issues = {}
        for col in df.columns:
            if column_types.get(col) == 'date_string':
                formats = set()
                sample = df[col].dropna().head(100)
                for val in sample:
                    val_str = str(val)
                    if re.match(r'\d{4}-\d{2}-\d{2}', val_str):
                        formats.add('YYYY-MM-DD')
                    elif re.match(r'\d{2}/\d{2}/\d{4}', val_str):
                        formats.add('DD/MM/YYYY')
                    elif re.match(r'\d{2}-\d{2}-\d{4}', val_str):
                        formats.add('DD-MM-YYYY')
                    elif re.match(r'\d{4}/\d{2}/\d{2}', val_str):
                        formats.add('YYYY/MM/DD')
                if len(formats) > 1:
                    date_issues[col] = list(formats)
        return date_issues

    @staticmethod
    def full_analysis(df):
        column_types = DataAnalyzer.detect_column_types(df)
        mixed_columns = DataAnalyzer.detect_mixed_columns(df)
        analysis = {
            'rows': int(len(df)),
            'columns': int(len(df.columns)),
            'column_names': list(df.columns),
            'column_types': column_types,
            'mixed_columns': mixed_columns,
            'missing_values': DataAnalyzer.analyze_missing_values(df),
            'duplicates': DataAnalyzer.detect_duplicates(df),
            'outliers': DataAnalyzer.detect_outliers(df, column_types),
            'text_issues': DataAnalyzer.analyze_text_issues(df, column_types),
            'date_formats': DataAnalyzer.detect_date_formats(df, column_types)
        }
        return convert_to_serializable(analysis)


# -------------------- NETTOYAGE DES DONNÉES --------------------

class DataCleaner:

    # ── 1. Doublons ───────────────────────────────────────────────────────────
    @staticmethod
    def remove_duplicates(df, uniqueness_threshold=0.95):
        initial_len = len(df)
        exact_count = int(df.duplicated().sum())
        df = df.drop_duplicates()
        candidate_cols = [
            col for col in df.columns
            if df[col].nunique(dropna=True) / len(df) < uniqueness_threshold
        ]
        structural_count = int(df.duplicated(subset=candidate_cols).sum()) if candidate_cols else 0
        if candidate_cols:
            df = df.drop_duplicates(subset=candidate_cols)
        return df, {
            "exact_duplicates_removed": exact_count,
            "structural_duplicates_removed": structural_count,
            "used_columns": candidate_cols,
            "total_removed": exact_count + structural_count,
            "initial_rows": initial_len,
            "final_rows": len(df)
        }

    # ── 2. Valeurs manquantes ─────────────────────────────────────────────────
    @staticmethod
    def handle_missing_values(df, column_types,
                               numeric_strategy: str = 'median',
                               text_strategy: str = 'mode'):
        """
        numeric_strategy : 'mean' | 'median' | 'mode' | 'zero' | 'drop'
        text_strategy    : 'mode' | 'empty' | 'drop'

        NOTE : les colonnes mixtes sont ignorées ici (gérées par handle_mixed_columns).
        """
        corrected = 0
        rows_to_drop = set()

        for col in df.columns:
            missing_mask = df[col].isna()
            count = int(missing_mask.sum())
            if count == 0:
                continue

            # ✅ Ignorer les colonnes mixtes — elles ont leur propre traitement
            if _is_truly_mixed_column(df[col]):
                continue

            col_type = column_types.get(col)

            if col_type == 'numeric':
                if numeric_strategy == 'mean':
                    df[col] = df[col].fillna(df[col].mean())
                    corrected += count
                elif numeric_strategy == 'median':
                    df[col] = df[col].fillna(df[col].median())
                    corrected += count
                elif numeric_strategy == 'mode':
                    mode_val = df[col].mode()
                    if len(mode_val):
                        df[col] = df[col].fillna(mode_val[0])
                        corrected += count
                elif numeric_strategy == 'zero':
                    df[col] = df[col].fillna(0)
                    corrected += count
                elif numeric_strategy == 'drop':
                    rows_to_drop.update(df[missing_mask].index.tolist())

            elif col_type in ('text', 'date_string'):
                if text_strategy == 'mode':
                    mode_val = df[col].mode()
                    if len(mode_val):
                        df[col] = df[col].fillna(mode_val[0])
                        corrected += count
                elif text_strategy == 'empty':
                    df[col] = df[col].fillna('')
                    corrected += count
                elif text_strategy == 'drop':
                    rows_to_drop.update(df[missing_mask].index.tolist())

        if rows_to_drop:
            df = df.drop(index=list(rows_to_drop)).reset_index(drop=True)

        return df, int(corrected)

    # ── 3. Colonnes mixtes ────────────────────────────────────────────────────
    @staticmethod
    def handle_mixed_columns(df, mixed_strategy: str = 'replace_median'):
        """
        Traite UNIQUEMENT les colonnes qui contiennent à la fois des valeurs
        numériques valides ET du texte non-numérique non-null.

        Les colonnes purement textuelles (ST_NAME = 'PUTNAM', 'BERKELEY'…)
        ne contiennent AUCUNE valeur numérique → _is_truly_mixed_column()
        retourne False → elles sont ignorées et leurs données préservées.

        mixed_strategy :
          'to_numeric'     → convertit en numérique (texte → NaN pour suite)
          'replace_median' → remplace les valeurs texte par la médiane
          'replace_mean'   → remplace les valeurs texte par la moyenne
          'drop_rows'      → supprime les lignes avec valeur texte inattendue
        """
        corrections = 0

        for col in df.columns:
            # ✅ GARDE-FOU CRITIQUE : ne traiter que les vraies colonnes mixtes
            if not _is_truly_mixed_column(df[col]):
                continue

            original_nulls = df[col].isna()
            numeric_attempt = pd.to_numeric(df[col], errors='coerce')
            new_nulls = numeric_attempt.isna()

            # Lignes non-nulles qui n'ont pas pu être converties en numérique
            mixed_mask = new_nulls & ~original_nulls

            if mixed_mask.sum() == 0:
                continue

            logging.debug(
                f"[MIXED] col={col}, strategy={mixed_strategy}, "
                f"n_mixed={int(mixed_mask.sum())}, "
                f"n_numeric={int((~new_nulls).sum())}"
            )

            if mixed_strategy == 'to_numeric':
                df[col] = numeric_attempt
                corrections += int(mixed_mask.sum())

            elif mixed_strategy == 'replace_median':
                median_val = numeric_attempt.median()
                df[col] = numeric_attempt
                df.loc[mixed_mask, col] = median_val
                corrections += int(mixed_mask.sum())

            elif mixed_strategy == 'replace_mean':
                mean_val = numeric_attempt.mean()
                df[col] = numeric_attempt
                df.loc[mixed_mask, col] = mean_val
                corrections += int(mixed_mask.sum())

            elif mixed_strategy == 'drop_rows':
                corrections += int(mixed_mask.sum())
                df = df[~mixed_mask].reset_index(drop=True)

        return df, int(corrections)

    # ── 4. Outliers ───────────────────────────────────────────────────────────
    @staticmethod
    def remove_outliers(df, column_types, method: str = 'median'):
        """method : 'median' | 'mean' | 'cap' | 'nan' | 'remove' | 'flag'"""
        initial = len(df)
        outliers_info = {}

        for col in df.columns:
            if column_types.get(col) != 'numeric':
                continue
            numeric_col = pd.to_numeric(df[col], errors='coerce')
            Q1 = numeric_col.quantile(0.25)
            Q3 = numeric_col.quantile(0.75)
            IQR = Q3 - Q1
            lower = Q1 - 1.5 * IQR
            upper = Q3 + 1.5 * IQR
            mask = (numeric_col < lower) | (numeric_col > upper)
            count = int(mask.sum())
            if count == 0:
                continue

            outliers_info[col] = count

            if method == 'remove':
                df = df[~mask]
            elif method == 'median':
                df.loc[mask, col] = numeric_col.median()
            elif method == 'mean':
                df.loc[mask, col] = numeric_col.mean()
            elif method == 'cap':
                df[col] = numeric_col.clip(lower=lower, upper=upper)
            elif method == 'nan':
                df.loc[mask, col] = np.nan
            elif method == 'flag':
                df[f'{col}_is_outlier'] = mask

        return df, {
            'outliers_detected': outliers_info,
            'rows_removed': int(initial - len(df)) if method == 'remove' else 0,
            'method_used': method,
            'total_outliers': sum(outliers_info.values()) if outliers_info else 0
        }

    # ── 5. Nettoyage texte ────────────────────────────────────────────────────
    @staticmethod
    def clean_text(df, column_types,
                   remove_emojis: bool = True,
                   remove_special_chars: bool = True,
                   trim_spaces: bool = True,
                   deduplicate_spaces: bool = True):
        corrections = 0
        emoji_pattern = re.compile(
            "["
            u"\U0001F600-\U0001F64F"
            u"\U0001F300-\U0001F5FF"
            u"\U0001F680-\U0001F6FF"
            u"\U0001F1E0-\U0001F1FF"
            "]+", flags=re.UNICODE
        )

        for col in df.columns:
            if column_types.get(col) not in ('text', 'date_string'):
                continue
            before = df[col].astype(str)
            s = df[col].apply(lambda x: '' if pd.isna(x) else str(x))
            if remove_emojis:
                s = s.apply(lambda x: emoji_pattern.sub('', x))
            if trim_spaces:
                s = s.str.strip()
            if deduplicate_spaces:
                s = s.apply(lambda x: re.sub(r'\s+', ' ', x))
            if remove_special_chars:
                s = s.apply(lambda x: re.sub(r'[^\w\s\-_.,;:!?\'@]', '', x))
            df[col] = s
            corrections += int((before != df[col].astype(str)).sum())

        return df, int(corrections)

    # ── 6. Harmonisation dates ────────────────────────────────────────────────
    @staticmethod
    def harmonize_dates(df, column_types, target_format: str = 'YYYY-MM-DD'):
        fmt_map = {
            'YYYY-MM-DD': '%Y-%m-%d',
            'DD/MM/YYYY': '%d/%m/%Y',
            'MM/DD/YYYY': '%m/%d/%Y',
        }
        strftime = fmt_map.get(target_format, '%Y-%m-%d')

        for col in df.columns:
            if column_types.get(col) == 'date_string':
                parsed = pd.to_datetime(df[col], errors='coerce', dayfirst=True)
                df[col] = parsed.apply(
                    lambda x: x.strftime(strftime) if pd.notna(x) else None
                )
        return df

    # ── 7. Normalisation casse ────────────────────────────────────────────────
    @staticmethod
    def normalize_case(df, column_types, case_style: str = 'title'):
        """case_style : 'title' | 'lower' | 'upper'"""
        corrections = 0
        for col in df.columns:
            if column_types.get(col) != 'text':
                continue
            before = df[col].apply(lambda x: '' if pd.isna(x) else str(x))
            if case_style == 'title':
                df[col] = df[col].apply(lambda x: '' if pd.isna(x) else str(x).title())
            elif case_style == 'lower':
                df[col] = df[col].apply(lambda x: '' if pd.isna(x) else str(x).lower())
            elif case_style == 'upper':
                df[col] = df[col].apply(lambda x: '' if pd.isna(x) else str(x).upper())
            after = df[col].apply(lambda x: '' if pd.isna(x) else str(x))
            corrections += int((before != after).sum())
        return df, int(corrections)

    # ── 8. Orchestrateur ──────────────────────────────────────────────────────
    @staticmethod
    def apply_cleaning(df, actions, column_types, outlier_method='median', options=None):
        """
        Ordre : doublons → mixtes → manquants → outliers → texte → dates → casse

        options (dict, clés optionnelles) :
          numericStrategy    'mean'|'median'|'mode'|'zero'|'drop'
          textStrategy       'mode'|'empty'|'drop'
          mixedStrategy      'to_numeric'|'replace_median'|'replace_mean'|'drop_rows'
          outlierMethod      'median'|'mean'|'cap'|'nan'|'remove'|'flag'
          removeEmojis       bool
          removeSpecialChars bool
          trimSpaces         bool
          deduplicateSpaces  bool
          targetFormat       'YYYY-MM-DD'|'DD/MM/YYYY'|'MM/DD/YYYY'
          caseStyle          'title'|'lower'|'upper'
        """
        if options is None:
            options = {}

        resolved_outlier_method = options.get(
            'outlierMethod', options.get('outlier_method', outlier_method)
        )

        results = {'initial_rows': int(len(df)), 'actions_performed': []}

        if 'duplicates' in actions:
            df, removed = DataCleaner.remove_duplicates(df)
            results['duplicates_removed'] = removed
            results['actions_performed'].append('Suppression des doublons')

        if 'mixed_columns' in actions:
            strategy = options.get('mixedStrategy', options.get('mixed_strategy', 'replace_median'))
            df, corrected = DataCleaner.handle_mixed_columns(df, mixed_strategy=strategy)
            results['mixed_columns_corrected'] = corrected
            results['actions_performed'].append(f'Colonnes mixtes ({strategy})')

        if 'missing_values' in actions:
            num_strat = options.get('numericStrategy', options.get('numeric_strategy', 'median'))
            txt_strat = options.get('textStrategy', options.get('text_strategy', 'mode'))
            df, corrected = DataCleaner.handle_missing_values(
                df, column_types,
                numeric_strategy=num_strat,
                text_strategy=txt_strat
            )
            results['missing_corrected'] = corrected
            results['actions_performed'].append(
                f'Valeurs manquantes (num: {num_strat}, txt: {txt_strat})'
            )

        if 'outliers' in actions:
            df, outlier_results = DataCleaner.remove_outliers(
                df, column_types, method=resolved_outlier_method
            )
            results['outliers_removed'] = outlier_results['rows_removed']
            results['outliers_info'] = outlier_results
            results['actions_performed'].append(f'Outliers ({resolved_outlier_method})')

        if 'text_cleaning' in actions:
            df, corrections = DataCleaner.clean_text(
                df, column_types,
                remove_emojis=options.get('removeEmojis', True),
                remove_special_chars=options.get('removeSpecialChars', True),
                trim_spaces=options.get('trimSpaces', True),
                deduplicate_spaces=options.get('deduplicateSpaces', True),
            )
            results['text_normalized'] = corrections
            results['actions_performed'].append('Nettoyage des textes')

        if 'date_format' in actions:
            target_fmt = options.get('targetFormat', options.get('target_format', 'YYYY-MM-DD'))
            df = DataCleaner.harmonize_dates(df, column_types, target_format=target_fmt)
            results['actions_performed'].append(f'Harmonisation des dates → {target_fmt}')

        if 'case_normalization' in actions:
            style = options.get('caseStyle', options.get('case_style', 'title'))
            df, corrections = DataCleaner.normalize_case(df, column_types, case_style=style)
            results['case_normalized'] = corrections
            results['actions_performed'].append(f'Normalisation de la casse ({style})')

        results['final_rows'] = int(len(df))
        results['columns'] = int(len(df.columns))
        return df, convert_to_serializable(results)


# -------------------- ROUTES API --------------------

@app.route('/api/upload', methods=['POST'])
@token_required
def upload_file(current_user_id):
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Aucun fichier fourni'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Nom de fichier vide'}), 400
        if not allowed_file(file.filename):
            return jsonify({'error': 'Format non supporté'}), 400

        session_id = str(uuid.uuid4())
        filename = secure_filename(file.filename)
        file_extension = filename.rsplit('.', 1)[1].lower()
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{session_id}_{filename}")
        file.save(filepath)

        df = load_file(filepath, file_extension)
        analysis = DataAnalyzer.full_analysis(df)

        sessions_db[session_id] = {
            'filename': filename,
            'file_extension': file_extension,
            'filepath': filepath,
            'dataframe': df,
            'analysis': analysis,
            'timestamp': datetime.now().isoformat(),
            'user_id': current_user_id,
            'operation_history': []      # ← Historique persistant par session
        }

        return jsonify({
            'session_id': session_id,
            'filename': filename,
            'analysis': analysis
        }), 200

    except Exception as e:
        logging.error(f"[UPLOAD ERROR] {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/sessions', methods=['GET'])
@token_required
def get_sessions(current_user_id):
    sessions_list = []
    for sid, session in sessions_db.items():
        if session.get('user_id') == current_user_id:
            sessions_list.append({
                'session_id': sid,
                'filename': session['filename'],
                'timestamp': session['timestamp'],
                'rows': session['analysis']['rows'],
                'columns': session['analysis']['columns'],
                'status': session.get('status', 'uploaded'),
                'operation_count': len(session.get('operation_history', []))
            })
    sessions_list.sort(key=lambda x: x['timestamp'], reverse=True)
    return jsonify({'sessions': sessions_list}), 200


@app.route('/api/session/<session_id>', methods=['GET'])
@token_required
def get_session_details(current_user_id, session_id):
    try:
        if session_id not in sessions_db:
            return jsonify({'error': 'Session non trouvée'}), 404

        session = sessions_db[session_id]

        if session.get('user_id') != current_user_id:
            return jsonify({'error': 'Accès non autorisé'}), 403

        response_data = {
            'session_id': session_id,
            'filename': session['filename'],
            'file_extension': session['file_extension'],
            'timestamp': session['timestamp'],
            'analysis': session['analysis'],
            'status': session.get('status', 'uploaded'),
            'operation_history': session.get('operation_history', [])
        }
        if 'cleaning_results' in session and 'cleaned_filepath' in session:
            response_data['cleaning_results'] = {
                'results': session['cleaning_results'],
                'cleaned_filepath': session['cleaned_filepath'],
                'cleaned_filename': session['cleaned_filename']
            }
            response_data['status'] = 'cleaned'

        return jsonify(response_data), 200

    except Exception as e:
        logging.error(f"[GET SESSION ERROR] {str(e)}")
        return jsonify({'error': str(e)}), 500


# ── Route dédiée : historique d'une session ───────────────────────────────────
@app.route('/api/session/<session_id>/history', methods=['GET'])
@token_required
def get_session_history(current_user_id, session_id):
    try:
        if session_id not in sessions_db:
            return jsonify({'error': 'Session non trouvée'}), 404

        session = sessions_db[session_id]

        if session.get('user_id') != current_user_id:
            return jsonify({'error': 'Accès non autorisé'}), 403

        return jsonify({
            'session_id': session_id,
            'filename': session['filename'],
            'timestamp': session['timestamp'],
            'status': session.get('status', 'uploaded'),
            'history': session.get('operation_history', [])
        }), 200

    except Exception as e:
        logging.error(f"[GET HISTORY ERROR] {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/clean', methods=['POST'])
@token_required
def clean_data(current_user_id):
    request_data = None
    session = None
    try:
        request_data = request.json
        session_id = request_data.get('session_id')
        actions = request_data.get('actions', [])
        outlier_method = request_data.get('outlier_method', 'median')
        options = request_data.get('options', {})

        if session_id not in sessions_db:
            return jsonify({'error': 'Session non trouvée'}), 404

        session = sessions_db[session_id]

        if session.get('user_id') != current_user_id:
            return jsonify({'error': 'Accès non autorisé'}), 403

        df = session['dataframe'].copy()
        column_types = session['analysis']['column_types']

        cleaned_df, results = DataCleaner.apply_cleaning(
            df, actions, column_types,
            outlier_method=outlier_method,
            options=options
        )

        cleaned_filename = f"cleaned_{session['filename']}"
        cleaned_filepath = os.path.join(
            app.config['CLEANED_FOLDER'], f"{session_id}_{cleaned_filename}"
        )
        save_file(cleaned_df, cleaned_filepath, session['file_extension'])

        session['cleaned_filepath'] = cleaned_filepath
        session['cleaned_filename'] = cleaned_filename
        session['cleaned_dataframe'] = cleaned_df
        session['cleaning_results'] = results
        session['status'] = 'cleaned'

        # ── Persister l'opération dans l'historique de la session ──────────
        history_entry = {
            'id': str(uuid.uuid4()),
            'timestamp': datetime.now().isoformat(),
            'actions': actions,
            'options': options,
            'results_summary': {
                'initial_rows': results.get('initial_rows'),
                'final_rows': results.get('final_rows'),
                'actions_performed': results.get('actions_performed', []),
                'duplicates_removed': (
                    results.get('duplicates_removed', {}).get('total_removed', 0)
                    if results.get('duplicates_removed') else 0
                ),
                'missing_corrected': results.get('missing_corrected', 0),
                'mixed_columns_corrected': results.get('mixed_columns_corrected', 0),
                'outliers_info': results.get('outliers_info'),
                'text_normalized': results.get('text_normalized', 0),
                'case_normalized': results.get('case_normalized', 0),
            },
            'status': 'success'
        }
        if 'operation_history' not in session:
            session['operation_history'] = []
        session['operation_history'].insert(0, history_entry)
        # ───────────────────────────────────────────────────────────────────

        return jsonify({
            'session_id': session_id,
            'results': results,
            'download_filename': cleaned_filename,
            'history_entry': history_entry
        }), 200

    except Exception as e:
        logging.error(f"[CLEAN ERROR] {str(e)}")
        # Enregistrer l'erreur dans l'historique
        if session is not None:
            if 'operation_history' not in session:
                session['operation_history'] = []
            session['operation_history'].insert(0, {
                'id': str(uuid.uuid4()),
                'timestamp': datetime.now().isoformat(),
                'actions': request_data.get('actions', []) if request_data else [],
                'options': request_data.get('options', {}) if request_data else {},
                'results_summary': {},
                'status': 'error',
                'error': str(e)
            })
        return jsonify({'error': str(e)}), 500


@app.route('/api/preview/<session_id>', methods=['GET'])
@token_required
def preview_data(current_user_id, session_id):
    try:
        if session_id not in sessions_db:
            return jsonify({'error': 'Session non trouvée'}), 404

        session = sessions_db[session_id]

        if session.get('user_id') != current_user_id:
            return jsonify({'error': 'Accès non autorisé'}), 403

        df = session['dataframe'].copy()
        df = df.replace({np.nan: None, np.inf: None, -np.inf: None})
        preview_df = df.head(100)

        rows = []
        for _, row in preview_df.iterrows():
            row_data = []
            for val in row:
                if val is None:
                    row_data.append(None)
                elif isinstance(val, (np.integer, np.int64)):
                    row_data.append(int(val))
                elif isinstance(val, (np.floating, np.float64)):
                    row_data.append(None if (np.isnan(val) or np.isinf(val)) else float(val))
                else:
                    try:
                        if pd.isna(val):
                            row_data.append(None)
                            continue
                    except Exception:
                        pass
                    row_data.append(str(val))
            rows.append(row_data)

        return jsonify({
            'columns': [str(col) for col in df.columns],
            'rows': rows,
            'total_rows': int(len(session['dataframe'])),
            'mixed_columns': session['analysis'].get('mixed_columns', [])
        }), 200

    except Exception as e:
        logging.error(f"[PREVIEW ERROR] {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/preview-cleaned/<session_id>', methods=['GET'])
@token_required
def preview_cleaned_data(current_user_id, session_id):
    try:
        if session_id not in sessions_db:
            return jsonify({'error': 'Session non trouvée'}), 404

        session = sessions_db[session_id]

        if session.get('user_id') != current_user_id:
            return jsonify({'error': 'Accès non autorisé'}), 403

        if 'cleaned_dataframe' not in session:
            return jsonify({'error': 'Données nettoyées non disponibles'}), 404

        df = session['cleaned_dataframe'].copy()
        df = df.replace({np.nan: None, np.inf: None, -np.inf: None})
        preview_df = df.head(100)

        rows = []
        for _, row in preview_df.iterrows():
            row_data = []
            for val in row:
                if val is None:
                    row_data.append(None)
                elif isinstance(val, (np.integer, np.int64)):
                    row_data.append(int(val))
                elif isinstance(val, (np.floating, np.float64)):
                    row_data.append(None if (np.isnan(val) or np.isinf(val)) else float(val))
                else:
                    try:
                        if pd.isna(val):
                            row_data.append(None)
                            continue
                    except Exception:
                        pass
                    row_data.append(str(val))
            rows.append(row_data)

        return jsonify({
            'columns': [str(col) for col in df.columns],
            'rows': rows,
            'total_rows': int(len(session['cleaned_dataframe']))
        }), 200

    except Exception as e:
        logging.error(f"[PREVIEW CLEANED ERROR] {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/download/<session_id>', methods=['GET'])
@token_required
def download_file(current_user_id, session_id):
    try:
        if session_id not in sessions_db:
            return jsonify({'error': 'Session non trouvée'}), 404

        session = sessions_db[session_id]

        if session.get('user_id') != current_user_id:
            return jsonify({'error': 'Accès non autorisé'}), 403

        if 'cleaned_filepath' not in session:
            return jsonify({'error': 'Fichier nettoyé non disponible'}), 404

        return send_file(
            session['cleaned_filepath'],
            as_attachment=True,
            download_name=session['cleaned_filename']
        )

    except Exception as e:
        logging.error(f"[DOWNLOAD ERROR] {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/download-multiple', methods=['POST'])
@token_required
def download_multiple(current_user_id):
    try:
        data = request.json
        session_ids = data.get('session_ids', [])

        if not session_ids:
            return jsonify({'error': 'Aucune session sélectionnée'}), 400
        if len(session_ids) > 10:
            return jsonify({'error': 'Maximum 10 fichiers à la fois'}), 400

        zip_buffer = BytesIO()
        with ZipFile(zip_buffer, 'w') as zip_file:
            for session_id in session_ids:
                if session_id not in sessions_db:
                    continue
                session = sessions_db[session_id]
                if session.get('user_id') != current_user_id:
                    continue
                if 'cleaned_filepath' not in session:
                    continue
                if not os.path.exists(session['cleaned_filepath']):
                    continue
                zip_file.write(session['cleaned_filepath'], arcname=session['cleaned_filename'])

        zip_buffer.seek(0)
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'data_cleaned_{datetime.now().strftime("%Y%m%d_%H%M%S")}.zip'
        )

    except Exception as e:
        logging.error(f"[DOWNLOAD MULTIPLE ERROR] {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/chat/recommend', methods=['POST'])
@token_required
def get_recommendations(current_user_id):
    try:
        data = request.json
        session_id = data.get('session_id')

        if session_id not in sessions_db:
            return jsonify({'error': 'Session non trouvée'}), 404

        session = sessions_db[session_id]

        if session.get('user_id') != current_user_id:
            return jsonify({'error': 'Accès non autorisé'}), 403

        assistant = DataAssistant(session['dataframe'], session['analysis'])
        recommendations = assistant.generate_recommendations()

        return jsonify({'recommendations': recommendations, 'count': len(recommendations)}), 200

    except Exception as e:
        logging.error(f"[RECOMMEND ERROR] {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/chat/ask', methods=['POST'])
@token_required
def ask_question(current_user_id):
    try:
        data = request.json
        session_id = data.get('session_id')
        question = data.get('question', '')

        if session_id not in sessions_db:
            return jsonify({'error': 'Session non trouvée'}), 404
        if not question:
            return jsonify({'error': 'Question vide'}), 400

        session = sessions_db[session_id]

        if session.get('user_id') != current_user_id:
            return jsonify({'error': 'Accès non autorisé'}), 403

        assistant = DataAssistant(session['dataframe'], session['analysis'])
        answer = assistant.answer_question(question)

        return jsonify(answer), 200

    except Exception as e:
        logging.error(f"[ASK ERROR] {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/chat/generate-report', methods=['POST'])
@token_required
def generate_report(current_user_id):
    try:
        data = request.json
        session_id = data.get('session_id')

        if session_id not in sessions_db:
            return jsonify({'error': 'Session non trouvée'}), 404

        session = sessions_db[session_id]

        if session.get('user_id') != current_user_id:
            return jsonify({'error': 'Accès non autorisé'}), 403

        assistant = DataAssistant(session['dataframe'], session['analysis'])
        recommendations = assistant.generate_recommendations()

        doc = Document()
        title = doc.add_heading('Rapport d\'Analyse de Qualité des Données', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('📊 Informations Générales', level=1)
        doc.add_paragraph(f"Fichier : {session['filename']}")
        doc.add_paragraph(f"Date d'analyse : {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Nombre de lignes : {session['analysis']['rows']:,}")
        doc.add_paragraph(f"Nombre de colonnes : {session['analysis']['columns']}")

        mixed_cols = session['analysis'].get('mixed_columns', [])
        if mixed_cols:
            doc.add_heading('⚠️ Colonnes Mixtes Détectées', level=1)
            doc.add_paragraph(
                f"{len(mixed_cols)} colonne(s) avec mélange numérique/texte : {', '.join(mixed_cols)}"
            )

        # Historique des opérations dans le rapport
        history = session.get('operation_history', [])
        if history:
            doc.add_heading('📋 Historique des Opérations', level=1)
            for op in history:
                ts = op.get('timestamp', '')
                try:
                    ts = datetime.fromisoformat(ts).strftime('%d/%m/%Y %H:%M:%S')
                except Exception:
                    pass
                status_label = '✅ Succès' if op.get('status') == 'success' else '❌ Erreur'
                doc.add_heading(f"Opération du {ts} — {status_label}", level=2)
                summary = op.get('results_summary', {})
                if summary:
                    doc.add_paragraph(
                        f"Lignes : {summary.get('initial_rows', '?')} → {summary.get('final_rows', '?')}"
                    )
                    for action in summary.get('actions_performed', []):
                        doc.add_paragraph(f"• {action}")
                if op.get('error'):
                    doc.add_paragraph(f"Erreur : {op['error']}")
                doc.add_paragraph("")

        dup = session['analysis'].get('duplicates', {})
        missing = session['analysis'].get('missing_values', {})
        total_dup = dup.get('exact_duplicates', 0) + dup.get('structural_duplicates', 0)
        total_missing = sum(v.get('count', 0) for v in missing.values())
        quality_score = assistant._calculate_quality_score(total_dup, total_missing)

        doc.add_heading('🎯 Score de Qualité', level=1)
        score_para = doc.add_paragraph(f"Score : {quality_score}/100")
        score_para.runs[0].font.size = Pt(16)
        score_para.runs[0].font.bold = True
        if quality_score >= 80:
            score_para.runs[0].font.color.rgb = RGBColor(34, 139, 34)
            doc.add_paragraph("✅ Excellente qualité")
        elif quality_score >= 60:
            score_para.runs[0].font.color.rgb = RGBColor(255, 165, 0)
            doc.add_paragraph("⚡ Qualité correcte")
        else:
            score_para.runs[0].font.color.rgb = RGBColor(220, 20, 60)
            doc.add_paragraph("⚠️ Nettoyage recommandé")

        doc.add_heading('💡 Recommandations', level=1)
        for i, rec in enumerate(recommendations, 1):
            doc.add_heading(f"{i}. {rec['title']}", level=2)
            doc.add_paragraph(f"Priorité : {rec['priority'].upper()}")
            doc.add_paragraph(f"Impact : {rec['impact']}")
            doc.add_paragraph(f"Justification : {rec['justification']}")
            doc.add_paragraph(f"Recommandé : {'✅ Oui' if rec['recommended'] else '⚠️ Optionnel'}")
            doc.add_paragraph("")

        report_filename = f"rapport_{session_id}.docx"
        report_path = os.path.join(app.config['CLEANED_FOLDER'], report_filename)
        doc.save(report_path)

        return jsonify({
            'message': 'Rapport généré avec succès',
            'download_url': f"/api/download-report/{session_id}",
            'filename': report_filename
        }), 200

    except Exception as e:
        logging.error(f"[GENERATE REPORT ERROR] {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/download-report/<session_id>', methods=['GET'])
@token_required
def download_report(current_user_id, session_id):
    try:
        if session_id in sessions_db:
            session = sessions_db[session_id]
            if session.get('user_id') != current_user_id:
                return jsonify({'error': 'Accès non autorisé'}), 403

        report_filename = f"rapport_{session_id}.docx"
        report_path = os.path.join(app.config['CLEANED_FOLDER'], report_filename)

        if not os.path.exists(report_path):
            return jsonify({'error': 'Rapport non trouvé'}), 404

        return send_file(
            report_path,
            as_attachment=True,
            download_name=f"rapport_analyse_{datetime.now().strftime('%Y%m%d')}.docx"
        )

    except Exception as e:
        logging.error(f"[DOWNLOAD REPORT ERROR] {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'healthy', 'timestamp': datetime.now().isoformat()}), 200


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port)